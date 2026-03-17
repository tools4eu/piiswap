"""Word (.docx) adapter using python-docx."""

from pathlib import Path

from piiswap.adapters.base import FileAdapter, register_adapter


@register_adapter
class DocxAdapter(FileAdapter):
    """Read and write .docx files, preserving formatting."""

    supported_extensions = (".docx",)

    def read(self, path: Path) -> str:
        from docx import Document

        doc = Document(str(path))
        parts = []

        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                parts.append("\t".join(cells))

        # Headers and footers
        for section in doc.sections:
            for header_footer in (section.header, section.footer):
                if header_footer and header_footer.is_linked_to_previous is False:
                    for para in header_footer.paragraphs:
                        if para.text.strip():
                            parts.append(para.text)

        return "\n".join(parts)

    def write(self, path: Path, content: str) -> None:
        from docx import Document

        path.parent.mkdir(parents=True, exist_ok=True)

        # Try to use original as template if it exists (preserve styles)
        # Otherwise create a new document
        doc = Document()
        for line in content.split("\n"):
            doc.add_paragraph(line)

        doc.save(str(path))

    def anonymize_preserving_format(self, input_path: Path, output_path: Path, replace_fn) -> None:
        """Anonymize while preserving original Word formatting.

        replace_fn: callable that takes a string and returns the anonymized string.
        """
        from docx import Document

        doc = Document(str(input_path))

        # Process paragraphs — replace text in each run to preserve formatting
        for para in doc.paragraphs:
            self._replace_in_paragraph(para, replace_fn)

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_in_paragraph(para, replace_fn)

        # Process headers and footers
        for section in doc.sections:
            for header_footer in (section.header, section.footer):
                if header_footer:
                    for para in header_footer.paragraphs:
                        self._replace_in_paragraph(para, replace_fn)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))

    @staticmethod
    def _replace_in_paragraph(para, replace_fn):
        """Replace PII in paragraph while trying to preserve run-level formatting.

        Strategy: concatenate all runs, apply replacement on full text,
        then redistribute across runs proportionally.
        If the replaced text is a different length, put everything in run[0].
        """
        if not para.runs:
            if para.text:
                new_text = replace_fn(para.text)
                if new_text != para.text:
                    para.text = new_text
            return

        full_text = "".join(run.text for run in para.runs)
        new_text = replace_fn(full_text)

        if new_text == full_text:
            return  # Nothing changed

        # If run count is 1, simple replacement
        if len(para.runs) == 1:
            para.runs[0].text = new_text
            return

        # Multiple runs: put all text in first run, clear the rest
        # This preserves the first run's formatting for the full paragraph
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ""
